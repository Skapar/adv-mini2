import PyPDF2
import docx
import spacy
import re

from pydantic import BaseModel
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from users.models import CustomUser
from django.core.exceptions import ValidationError

from .schemas import ResumeAnalysis

nlp = spacy.load("en_core_web_sm")

TRENDING_SKILLS = {
    'tech': ['python', 'javascript', 'react', 'aws', 'docker', 'sql', 'machine learning', 'typescript', 'kubernetes',
             'cloud'],
    'soft': ['communication', 'leadership', 'problem-solving', 'adaptability', 'collaboration']
}

ATS_KEYWORDS = [
    'experience', 'development', 'software', 'programming', 'team', 'project', 'management',
    'skills', 'technology', 'design', 'implementation', 'optimization', 'data', 'analysis'
]

def extract_text_from_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            print(f"Extracted text from PDF: {text[:100]}...")
            return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        print(f"Extracted text from DOCX: {text[:100]}...")
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""
    
class ResumeAnalysis(BaseModel):
    skills: str
    experience: str
    education: str
    rating: float
    feedback: dict


def process_resume(file_path: str) -> dict:
    doc = docx.Document(file_path)
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    print(f"Extracted text from DOCX: {full_text[:100]}...")

    processed_text = full_text.lower()
    print(f"Processed text: {processed_text[:100]}...")

    skills_section = re.search(r'skills\s*([\s\S]*?)(work experience|education|$)', processed_text, re.IGNORECASE)
    skills = set()
    if skills_section:
        skills_text = skills_section.group(1).strip()
        skills = {skill.strip() for skill in skills_text.split('\n') if
                  skill.strip() and not skill.strip().startswith(('email', 'phone', 'name'))}

        known_skills = {'python', 'javascript', 'react', 'aws', 'docker', 'figma', 'adobe creative suite',
                        'communication', 'leadership'}
        skills = {s for s in skills if s in known_skills or len(s) > 2}
    skills_str = ', '.join(skills) if skills else 'unknown'
    print(f"Extracted skills: {skills}")


    experience_section = re.search(r'(work experience|experience)\s*([\s\S]*?)(education|$)', processed_text,
                                   re.IGNORECASE)
    experience_years = 0.0
    if experience_section:
        experience_text = experience_section.group(2)
        years = re.findall(r'(\d{4})\s*[–-]\s*(\d{4}|present)', experience_text, re.IGNORECASE)
        for start, end in years:
            end_year = 2025 if end.lower() == 'present' else int(end)
            experience_years += end_year - int(start)
    experience_str = f"{experience_years:.1f} years"
    print(f"Calculated experience: {experience_str}")


    education_section = re.search(r'education\s*([\s\S]*?)(languages|$)', processed_text, re.IGNORECASE)
    education = education_section.group(1).strip() if education_section else 'Not specified'
    print(f"Extracted education: {education}")

    rating = 20 + (len(skills) * 10 if skills else 0) + (experience_years * 5)
    rating = min(rating, 100)
    print(f"Calculated rating: {rating}")


    trending_tech_skills = {'python', 'javascript', 'react', 'aws', 'docker'}
    trending_soft_skills = {'communication', 'leadership', 'problem-solving'}
    skill_gaps = [
        f"Missing trending technical skills: {', '.join(trending_tech_skills - skills)}" if not skills & trending_tech_skills else "",
        f"Missing trending soft skills: {', '.join(trending_soft_skills - skills)}" if not skills & trending_soft_skills else ""
    ]

    formatting_tips = []
    if len(full_text.split('\n')) < 5:
        formatting_tips.append("Add more sections like Experience and Education for clarity.")
    if not skills_section:
        formatting_tips.append("Include a clear 'Skills' section.")

    ats_keywords = []
    if not skills or len(skills) < 3:
        ats_keywords.append("Add ATS-friendly keywords: development, software, programming, teamwork.")

    feedback = {
        'skill_gaps': [gap for gap in skill_gaps if gap],
        'formatting': formatting_tips,
        'ats_keywords': ats_keywords
    }
    print(f"Generated feedback: {feedback}")

    try:
        analysis = ResumeAnalysis(
            skills=skills_str,
            experience=experience_str,
            education=education,
            rating=rating,
            feedback=feedback
        )
        return analysis.dict()
    except ValidationError as e:
        print(f"Validation error: {e}")
        raise

def process_job_description(description):

    doc = nlp(description.lower())

    skills_keywords = [
        'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'django',
        'communication', 'teamwork', 'leadership', 'management', 'excel', 'aws',
        'docker', 'git', 'linux', 'agile', 'scrum', 'typescript', 'kubernetes',
        'cloud', 'machine learning', 'problem-solving', 'adaptability', 'collaboration'
    ]
    required_skills = set()
    for token in doc:
        if token.text in skills_keywords:
            required_skills.add(token.text)


    experience_pattern = r'(\d+)\s*(?:year|yr|month|experience)'
    experience_matches = re.findall(experience_pattern, description.lower())
    required_experience = 0.0
    for match in experience_matches:
        if match.isdigit():
            num = int(match)
            if num < 12 and 'month' in description.lower():
                required_experience += num / 12
            else:
                required_experience += num

    return {
        'required_skills': required_skills,
        'required_experience': required_experience,
        'text': description
    }


def match_resume_to_job(resume, job):
    resume_skills = set(resume.skills.split(', ')) if resume.skills else set()
    job_skills = set(job.required_skills.split(', ')) if job.required_skills else set()

    resume_experience_str = resume.experience or "0 years"
    resume_experience = float(re.search(r'(\d+\.\d+|\d+)', resume_experience_str).group()) if re.search(
        r'(\d+\.\d+|\d+)', resume_experience_str) else 0.0

    job_experience_str = str(
        job.required_experience) if job.required_experience is not None else "0 years"
    job_experience = float(re.search(r'(\d+\.\d+|\d+)', job_experience_str).group()) if re.search(r'(\d+\.\d+|\d+)',
                                                                                                  job_experience_str) else 0.0

    common_skills = resume_skills.intersection(job_skills)
    skills_match = len(common_skills) / len(job_skills) if job_skills else 0
    skills_score = min(skills_match * 50, 50)

    experience_diff = resume_experience - job_experience
    if experience_diff >= 0:
        experience_score = min((resume_experience / job_experience) * 30, 30) if job_experience else 30
    else:
        experience_score = max(0, 30 + experience_diff * 5)

    try:
        tfidf = TfidfVectorizer(stop_words='english')
        resume_text = f"{resume.skills or ''} {resume.experience or ''} {resume.education or ''}"
        texts = [resume_text, job.description]
        tfidf_matrix = tfidf.fit_transform(texts)
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        text_score = similarity * 20
    except Exception as e:
        print(f"TF-IDF error for resume {resume.id}: {str(e)}")
        text_score = 0

    compatibility_score = skills_score + experience_score + text_score

    try:
        user = CustomUser.objects.get(id=resume.user_id)
        username = user.username
    except CustomUser.DoesNotExist:
        username = "Unknown"

    return {
        'resume_id': str(resume.id),
        'user': username,
        'compatibility_score': round(compatibility_score, 2),
        'matched_skills': ', '.join(common_skills),
        'resume_skills': resume.skills,
        'resume_experience': resume.experience
    }